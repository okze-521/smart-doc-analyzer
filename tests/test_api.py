"""RAG API 集成测试"""

import io
import pytest
from pathlib import Path


def _make_sample_docx_bytes():
    """返回 DOCX 字节"""
    from docx import Document
    doc = Document()
    doc.add_heading("测试文档标题", level=1)
    doc.add_paragraph("这是正文内容。")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_sample_xlsx_bytes():
    """返回 XLSX 字节"""
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "产品"
    ws["B1"] = "销量"
    ws["A2"] = "手机"
    ws["B2"] = 1500
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


class TestRagAPI:
    """RAG API 基本端点 — 使用 conftest 的 client fixture"""

    def test_health(self, client):
        r = client.get("/health")
        assert r.status_code == 200
        assert r.json()["status"] == "ok"

    def test_upload_no_file(self, client):
        r = client.post("/api/v1/documents/upload")
        assert r.status_code == 422

    def test_upload_wrong_type(self, client):
        r = client.post(
            "/api/v1/documents/upload",
            files={"file": ("test.txt", b"hello", "text/plain")},
        )
        assert r.status_code == 400

    def test_search_empty_query(self, client):
        r = client.post("/api/v1/documents/search", json={"query": ""})
        assert r.status_code == 422

    def test_search_minimal(self, client):
        r = client.post("/api/v1/documents/search", json={"query": "测试"})
        assert r.status_code == 200
        data = r.json()
        assert "chunks" in data
        assert "total_hits" in data

    def test_list_documents_empty(self, client):
        r = client.get("/api/v1/documents")
        assert r.status_code == 200
        data = r.json()
        assert data["items"] == []

    def test_get_not_found(self, client):
        r = client.get("/api/v1/documents/99999")
        assert r.status_code == 404

    def test_upload_pdf(self, client):
        from tests.conftest import _make_minimal_pdf_bytes
        pdf_bytes = _make_minimal_pdf_bytes()

        r = client.post(
            "/api/v1/documents/upload",
            files={"file": ("test.pdf", pdf_bytes, "application/pdf")},
        )
        assert r.status_code == 200
        data = r.json()
        assert data["file_type"] == "pdf"
        assert data["chunk_count"] > 0
        assert data["status"] == "completed"

    def test_upload_then_search(self, client):
        from tests.conftest import _make_minimal_pdf_bytes

        pdf_bytes = _make_minimal_pdf_bytes()
        r = client.post(
            "/api/v1/documents/upload",
            files={"file": ("doc.pdf", pdf_bytes, "application/pdf")},
        )
        assert r.status_code == 200

        r = client.post("/api/v1/documents/search", json={"query": "Hello", "top_k": 3})
        assert r.status_code == 200
        data = r.json()
        assert data["total_hits"] > 0

    def test_upload_docx(self, client):
        r = client.post(
            "/api/v1/documents/upload",
            files={"file": ("doc.docx", _make_sample_docx_bytes(),
                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert r.status_code == 200
        data = r.json()
        assert data["file_type"] == "docx"
        assert data["chunk_count"] > 0

    def test_upload_xlsx(self, client):
        r = client.post(
            "/api/v1/documents/upload",
            files={"file": ("sheet.xlsx", _make_sample_xlsx_bytes(),
                   "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
        )
        assert r.status_code == 200
        data = r.json()
        assert data["file_type"] == "xlsx"
        assert data["chunk_count"] > 0
