"""pytest 全局 fixtures — 测试数据库 + TestClient + 测试文档"""

import io
from pathlib import Path

import pytest
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from src.database import Base, get_db

# 测试用 SQLite 内存数据库（每次测试自动创建/销毁）
TEST_DATABASE_URL = "sqlite:///./data/test.db"


@pytest.fixture
def db_session():
    """每个测试独立的数据库 session"""
    engine = create_engine(
        TEST_DATABASE_URL,
        connect_args={"check_same_thread": False},
    )
    Base.metadata.create_all(bind=engine)
    TestingSession = sessionmaker(bind=engine)
    session = TestingSession()
    try:
        yield session
    finally:
        session.close()
        Base.metadata.drop_all(bind=engine)


@pytest.fixture
def client(db_session):
    """FastAPI TestClient — 用于测试 API 端点"""
    from src.main import app
    from src.api.rag import get_ingest_service, get_search_service
    from src.core.ingest_service import IngestService
    from src.core.search_service import SearchService
    from src.core.embedder import TextEmbedder

    # 使用本地 ModelScope 模型，避免网络下载
    MODEL_PATH = "models/models/AI-ModelScope--bge-small-zh-v1.5/snapshots/master"
    embedder = TextEmbedder(model_name=MODEL_PATH)

    def override_get_db():
        try:
            yield db_session
        finally:
            pass

    def override_ingest():
        return IngestService(embedder=embedder)

    def override_search():
        return SearchService(embedder=embedder)

    # 确保 Qdrant 集合存在
    from src.core.qdrant_store import QdrantStore
    qst = QdrantStore()
    qst.ensure_collection(vector_size=512)  # BGE-small-zh-v1.5 = 512 维

    app.dependency_overrides[get_db] = override_get_db
    app.dependency_overrides[get_ingest_service] = override_ingest
    app.dependency_overrides[get_search_service] = override_search

    with TestClient(app) as c:
        yield c
    app.dependency_overrides.clear()


# ============================================================
#  测试文档 Fixtures
# ============================================================

def _make_minimal_pdf_bytes() -> bytes:
    """返回最小有效 PDF 的字节"""
    # --- 构建 PDF 对象 ---
    obj1 = b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
    obj2 = b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
    obj3 = (
        b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R"
        b"/Resources<</Font<</F1<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>>>>>"
        b"/Contents 4 0 R>>endobj\n"
    )
    stream_data = (
        b"BT\n"
        b"/F1 14 Tf\n"
        b"50 700 Td (Hello PDF World) Tj\n"
        b"0 -20 Td (This is a test PDF document.) Tj\n"
        b"0 -20 Td (It contains multiple lines of text.) Tj\n"
        b"ET"
    )
    obj4 = b"4 0 obj<</Length %d>>stream\n" % len(stream_data) + stream_data + b"\nendstream\nendobj\n"

    # --- 构建 PDF 文件，记录 xref 偏移 ---
    header = b"%%PDF-1.4\n"
    offsets: list[int] = []

    body = bytearray(header)
    for obj in [obj1, obj2, obj3, obj4]:
        offsets.append(len(body))
        body.extend(obj)

    xref_offset = len(body)
    n_objects = len(offsets)
    xref = f"xref\n0 {n_objects + 1}\n0000000000 65535 f \n".encode("ascii")
    for off in offsets:
        xref += f"{off:010d} 00000 n \n".encode("ascii")

    trailer = f"trailer<</Size {n_objects + 1}/Root 1 0 R>>\nstartxref\n{xref_offset}\n%%EOF".encode("ascii")

    body.extend(xref)
    body.extend(trailer)

    return bytes(body)


def _make_minimal_pdf(filepath: Path) -> Path:
    """创建最小有效 PDF，动态计算 xref 偏移量"""
    filepath.write_bytes(_make_minimal_pdf_bytes())
    return filepath


@pytest.fixture
def sample_pdf(tmp_path):
    """创建测试用 PDF 文件"""
    return _make_minimal_pdf(tmp_path / "test.pdf")


@pytest.fixture
def sample_docx(tmp_path):
    """创建测试用 DOCX 文件（含标题、段落、表格）"""
    from docx import Document

    doc = Document()
    doc.add_heading("测试文档标题", level=1)
    doc.add_paragraph("这是第一段正文内容，用于测试解析器。")
    doc.add_heading("第二章", level=2)
    doc.add_paragraph("第二段内容，包含更多文字。")
    # 添加表格
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "年龄"
    table.cell(1, 0).text = "张三"
    table.cell(1, 1).text = "28"
    table.cell(2, 0).text = "李四"
    table.cell(2, 1).text = "32"

    filepath = tmp_path / "test.docx"
    doc.save(str(filepath))
    return filepath


@pytest.fixture
def sample_xlsx(tmp_path):
    """创建测试用 XLSX 文件"""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws["A1"] = "产品名称"
    ws["B1"] = "销量"
    ws["C1"] = "单价"
    ws["A2"] = "手机"
    ws["B2"] = 1500
    ws["C2"] = 2999
    ws["A3"] = "电脑"
    ws["B3"] = 800
    ws["C3"] = 5999

    # 第二个工作表
    ws2 = wb.create_sheet("汇总")
    ws2["A1"] = "总销量"
    ws2["B1"] = 2300

    filepath = tmp_path / "test.xlsx"
    wb.save(str(filepath))
    return filepath


@pytest.fixture
def sample_txt(tmp_path):
    """创建测试用纯文本文件"""
    filepath = tmp_path / "test.txt"
    filepath.write_text("这是纯文本测试文件。\n用于验证不支持的格式处理。", encoding="utf-8")
    return filepath
