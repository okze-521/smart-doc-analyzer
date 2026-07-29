"""文档解析引擎 — 策略模式，支持 PDF / DOCX / XLSX"""

from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class ParsedSection:
    """解析后的文档片段"""

    section_type: str = "paragraph"  # heading / paragraph / table
    level: int = 0  # 标题级别 (h1=1, h2=2, ...)
    content: str = ""  # 文本内容
    table_data: list = field(default_factory=list)  # 表格数据 [[row], [row]]
    page_number: int = 1
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """完整解析结果"""

    filename: str
    file_type: str
    total_pages: int
    sections: list[ParsedSection]
    metadata: dict
    full_text: str


class DocumentParser:
    """文档解析引擎 — 根据扩展名自动选择解析策略"""

    # ----------------------------------------------------------------
    #  入口
    # ----------------------------------------------------------------

    def parse(self, file_path: Path) -> ParsedDocument:
        """解析文档，自动识别格式"""
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = file_path.suffix.lower()
        if ext == ".pdf":
            return self._parse_pdf(file_path)
        elif ext == ".docx":
            return self._parse_docx(file_path)
        elif ext == ".xlsx":
            return self._parse_xlsx(file_path)
        else:
            raise ValueError(f"不支持的格式: {ext}（支持 .pdf / .docx / .xlsx）")

    # ----------------------------------------------------------------
    #  PDF 解析
    # ----------------------------------------------------------------

    def _parse_pdf(self, file_path: Path) -> ParsedDocument:
        import fitz  # pymupdf

        doc = fitz.open(str(file_path))
        total_pages = len(doc)
        all_text: list[str] = []
        sections: list[ParsedSection] = []

        for page_num, page in enumerate(doc, start=1):
            text = page.get_text()
            if text.strip():
                all_text.append(text)
                sections.append(
                    ParsedSection(
                        section_type="paragraph",
                        content=text.strip(),
                        page_number=page_num,
                    )
                )

        # 元数据
        meta = doc.metadata or {}
        metadata = {
            "author": str(meta.get("author", "")),
            "title": str(meta.get("title", "")),
            "creator": str(meta.get("creator", "")),
            "pages": total_pages,
        }

        return ParsedDocument(
            filename=file_path.name,
            file_type="pdf",
            total_pages=total_pages,
            sections=sections,
            metadata=metadata,
            full_text="\n".join(all_text),
        )

    # ----------------------------------------------------------------
    #  DOCX 解析
    # ----------------------------------------------------------------

    def _parse_docx(self, file_path: Path) -> ParsedDocument:
        from docx import Document

        doc = Document(str(file_path))
        sections: list[ParsedSection] = []
        all_text: list[str] = []

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # 段落
                para = _find_paragraph(doc, element)
                if para is None:
                    continue
                text = para.text.strip()
                if not text:
                    continue

                # 判断是否是标题
                if para.style.name.startswith("Heading") or para.style.name.startswith("标题"):
                    level_str = para.style.name.replace("Heading", "").replace("标题", "").strip()
                    try:
                        level = int(level_str)
                    except ValueError:
                        level = 1
                    sections.append(
                        ParsedSection(section_type="heading", level=level, content=text)
                    )
                else:
                    sections.append(ParsedSection(section_type="paragraph", content=text))
                all_text.append(text)

            elif tag == "tbl":
                # 表格
                table = _find_table(doc, element)
                if table is None:
                    continue
                rows = [[cell.text for cell in row.cells] for row in table.rows]
                sections.append(
                    ParsedSection(section_type="table", table_data=rows, content=_table_to_text(rows))
                )
                all_text.append(_table_to_text(rows))

        return ParsedDocument(
            filename=file_path.name,
            file_type="docx",
            total_pages=1,  # DOCX 无固定页数
            sections=sections,
            metadata={"paragraphs": len([s for s in sections if s.section_type == "paragraph"])},
            full_text="\n".join(all_text),
        )

    # ----------------------------------------------------------------
    #  XLSX 解析
    # ----------------------------------------------------------------

    def _parse_xlsx(self, file_path: Path) -> ParsedDocument:
        from openpyxl import load_workbook

        wb = load_workbook(str(file_path), read_only=True, data_only=True)
        sections: list[ParsedSection] = []
        all_text: list[str] = []
        sheet_names: list[str] = []

        for sheet_name in wb.sheetnames:
            sheet_names.append(sheet_name)
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                # 跳过全空行
                if any(cell is not None for cell in row):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    rows.append(row_values)

            if rows:
                sections.append(
                    ParsedSection(
                        section_type="table",
                        table_data=rows,
                        content=_table_to_text(rows),
                        metadata={"sheet_name": sheet_name},
                    )
                )
                all_text.append(f"[工作表: {sheet_name}]\n{_table_to_text(rows)}")

        wb.close()

        return ParsedDocument(
            filename=file_path.name,
            file_type="xlsx",
            total_pages=len(sheet_names),
            sections=sections,
            metadata={"sheet_names": sheet_names, "sheet_count": len(sheet_names)},
            full_text="\n\n".join(all_text),
        )


# ================================================================
#  辅助函数
# ================================================================

def _find_paragraph(doc, xml_element):
    """通过 XML 元素找到对应的 python-docx Paragraph 对象"""
    for para in doc.paragraphs:
        if para._element is xml_element:
            return para
    return None


def _find_table(doc, xml_element):
    """通过 XML 元素找到对应的 python-docx Table 对象"""
    for table in doc.tables:
        if table._element is xml_element:
            return table
    return None


def _table_to_text(rows: list[list[str]]) -> str:
    """表格转可读文本，供 full_text 使用"""
    if not rows:
        return ""
    lines = [" | ".join(row) for row in rows]
    return "\n".join(lines)
